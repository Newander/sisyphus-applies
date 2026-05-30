import base64
import ctypes
import json
import logging
import sys
import zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree

from backend.schemas import DocumentItem, DocumentPreview

DOCUMENT_EXTENSIONS = {".pdf", ".doc", ".docx", ".rtf", ".txt", ".md"}
TEXT_PREVIEW_EXTENSIONS = {".md", ".rtf", ".txt"}
DOCX_PREVIEW_EXTENSIONS = {".docx"}
CREATABLE_TEXT_EXTENSIONS = {".txt", ".md", ".docx"}
PREVIEW_LINES_LIMIT = 50
logger = logging.getLogger(__name__)


def encode_document_id(relative_path: str) -> str:
    raw = relative_path.encode("utf-8")
    return base64.urlsafe_b64encode(raw).decode("ascii").rstrip("=")


def decode_document_id(document_id: str) -> str:
    padding = "=" * (-len(document_id) % 4)
    return base64.urlsafe_b64decode(f"{document_id}{padding}").decode("utf-8")


def infer_document_type(path: Path) -> str:
    normalized = path.name.lower()
    if "cover" in normalized or "letter" in normalized:
        return "cover_letter"
    if "cv" in normalized or "resume" in normalized:
        return "cv"
    return "other"


def document_metadata_path(path: Path) -> Path:
    return path.with_name(f"{path.name}.meta.json")


def read_document_metadata(path: Path) -> dict[str, object]:
    metadata_path = document_metadata_path(path)
    if not metadata_path.is_file():
        return {}

    try:
        with metadata_path.open("r", encoding="utf-8") as file:
            payload = json.load(file)
    except (OSError, json.JSONDecodeError):
        logger.warning("Failed to read document metadata path=%s", metadata_path)
        return {}

    if not isinstance(payload, dict):
        return {}
    return payload


def write_document_metadata(
    path: Path,
    *,
    document_type: str,
    company_id: int | None,
    company_name: str | None,
) -> None:
    payload = {
        "document_type": document_type,
        "company_id": company_id,
        "company_name": company_name,
    }
    with document_metadata_path(path).open("w", encoding="utf-8") as file:
        json.dump(payload, file, ensure_ascii=False, indent=2)


def to_document_item(item: Path, storage_dir: Path) -> DocumentItem:
    relative_path = item.relative_to(storage_dir).as_posix()
    stat = item.stat()
    metadata = read_document_metadata(item)
    metadata_document_type = metadata.get("document_type")
    metadata_company_id = metadata.get("company_id")
    metadata_company_name = metadata.get("company_name")
    return DocumentItem(
        id=encode_document_id(relative_path),
        name=item.name,
        path=str(item),
        size_bytes=stat.st_size,
        modified_at=datetime.fromtimestamp(stat.st_mtime).astimezone(),
        document_type=(
            metadata_document_type
            if isinstance(metadata_document_type, str)
            else infer_document_type(item)
        ),
        company_id=metadata_company_id if isinstance(metadata_company_id, int) else None,
        company_name=metadata_company_name if isinstance(metadata_company_name, str) else None,
    )


def list_storage_documents(storage_dir: Path, limit: int | None = 20) -> list[DocumentItem]:
    storage_dir.mkdir(parents=True, exist_ok=True)
    storage_dir = storage_dir.resolve()
    logger.debug("Scanning documents storage_dir=%s limit=%s", storage_dir, limit)

    files = [
        item
        for item in storage_dir.iterdir()
        if item.is_file() and item.suffix.lower() in DOCUMENT_EXTENSIONS
    ]
    files.sort(key=lambda item: item.stat().st_mtime, reverse=True)
    selected_files = files if limit is None else files[:limit]

    documents = [to_document_item(item, storage_dir) for item in selected_files]
    logger.debug("Scanned documents total_files=%s returned=%s", len(files), len(documents))
    return documents


def list_storage_documents_page(
    storage_dir: Path,
    *,
    page: int,
    page_size: int,
    sort: str,
    direction: str,
) -> tuple[list[DocumentItem], int]:
    documents = list_storage_documents(storage_dir, limit=None)
    sort_getters = {
        "company_name": lambda item: item.company_name or "",
        "document_type": lambda item: item.document_type,
        "modified_at": lambda item: item.modified_at,
        "name": lambda item: item.name,
        "size_bytes": lambda item: item.size_bytes,
    }
    sort_key = sort_getters.get(sort, sort_getters["modified_at"])
    documents.sort(key=sort_key, reverse=direction == "desc")
    offset = (page - 1) * page_size
    return documents[offset : offset + page_size], len(documents)


def resolve_storage_document_path(storage_dir: Path, document_id: str) -> Path | None:
    storage_dir.mkdir(parents=True, exist_ok=True)
    storage_dir = storage_dir.resolve()
    try:
        relative_path = decode_document_id(document_id)
    except (ValueError, UnicodeDecodeError):
        logger.warning("Invalid document id document_id=%s", document_id)
        return None

    candidate = (storage_dir / relative_path).resolve()
    logger.debug("Resolved document document_id=%s candidate=%s", document_id, candidate)

    if not candidate.is_relative_to(storage_dir):
        logger.warning(
            "Document path escapes storage document_id=%s candidate=%s",
            document_id,
            candidate,
        )
        return None
    if not candidate.is_file() or candidate.suffix.lower() not in DOCUMENT_EXTENSIONS:
        logger.warning(
            "Document candidate is unavailable document_id=%s candidate=%s",
            document_id,
            candidate,
        )
        return None

    return candidate


def get_storage_document(storage_dir: Path, document_id: str) -> DocumentItem | None:
    storage_dir = storage_dir.resolve()
    candidate = resolve_storage_document_path(storage_dir, document_id)
    if candidate is None:
        return None
    return to_document_item(candidate, storage_dir)


def write_docx(text: str, path: Path) -> None:
    from docx import Document

    doc = Document()
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped:
            doc.add_paragraph(stripped)
        else:
            doc.add_paragraph("")
    doc.save(path)


def safe_text_document_name(file_name: str) -> str:
    name = Path(file_name.strip()).name.strip().strip(".")
    if not name:
        raise ValueError("File name is required")

    for character in '<>:"/\\|?*':
        name = name.replace(character, "_")
    name = " ".join(name.split())
    if not name:
        raise ValueError("File name is required")

    path = Path(name)
    if path.suffix.lower() not in CREATABLE_TEXT_EXTENSIONS:
        name = f"{name}.txt"
    return name


def create_text_document(
    storage_dir: Path,
    *,
    file_name: str,
    text: str,
    document_type: str,
    company_id: int | None,
    company_name: str | None,
) -> DocumentItem:
    storage_dir.mkdir(parents=True, exist_ok=True)
    storage_dir = storage_dir.resolve()
    safe_name = safe_text_document_name(file_name)
    document_path = (storage_dir / safe_name).resolve()

    if not document_path.is_relative_to(storage_dir):
        raise ValueError("File name is invalid")
    if document_path.exists():
        raise FileExistsError("Document with this file name already exists")

    if document_path.suffix.lower() == ".docx":
        write_docx(text, document_path)
    else:
        document_path.write_text(text, encoding="utf-8")
    write_document_metadata(
        document_path,
        document_type=document_type,
        company_id=company_id,
        company_name=company_name,
    )
    return to_document_item(document_path, storage_dir)


def _limited_preview(lines: list[str], line_limit: int) -> DocumentPreview:
    limited_lines = lines[:line_limit]
    return DocumentPreview(
        lines=limited_lines,
        line_count=len(limited_lines),
        truncated=len(lines) > line_limit,
    )


def _read_docx_preview(path: Path, line_limit: int) -> DocumentPreview:
    try:
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read("word/document.xml")
    except (KeyError, zipfile.BadZipFile, OSError):
        return DocumentPreview(
            lines=[],
            line_count=0,
            truncated=False,
            unsupported_reason="Failed to read DOCX file.",
        )

    try:
        root = ElementTree.fromstring(document_xml)
    except ElementTree.ParseError:
        return DocumentPreview(
            lines=[],
            line_count=0,
            truncated=False,
            unsupported_reason="Failed to parse DOCX file text.",
        )

    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    lines = []
    for paragraph in root.iter(f"{namespace}p"):
        text = "".join(text_node.text or "" for text_node in paragraph.iter(f"{namespace}t"))
        if text:
            lines.append(text)

    return _limited_preview(lines, line_limit)


def read_document_preview(path: Path, line_limit: int = PREVIEW_LINES_LIMIT) -> DocumentPreview:
    suffix = path.suffix.lower()
    if suffix in DOCX_PREVIEW_EXTENSIONS:
        return _read_docx_preview(path, line_limit)

    if suffix not in TEXT_PREVIEW_EXTENSIONS:
        return DocumentPreview(
            lines=[],
            line_count=0,
            truncated=False,
            unsupported_reason="Preview is available only for TXT, MD, RTF and DOCX files.",
        )

    for encoding in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            with path.open("r", encoding=encoding) as file:
                lines = []
                for index, line in enumerate(file):
                    if index >= line_limit:
                        return DocumentPreview(
                            lines=lines,
                            line_count=len(lines),
                            truncated=True,
                        )
                    lines.append(line.rstrip("\r\n"))
            return DocumentPreview(lines=lines, line_count=len(lines), truncated=False)
        except UnicodeDecodeError:
            continue

    return DocumentPreview(
        lines=[],
        line_count=0,
        truncated=False,
        unsupported_reason="Failed to read file as text.",
    )


def send_document_to_recycle_bin(path: Path) -> None:
    metadata_path = document_metadata_path(path)

    if sys.platform == "win32":
        _windows_recycle(path)
    else:
        path.unlink()

    if metadata_path.exists():
        metadata_path.unlink()


def _windows_recycle(path: Path) -> None:
    from ctypes import wintypes

    class SHFILEOPSTRUCTW(ctypes.Structure):
        _fields_ = [
            ("hwnd", wintypes.HWND),
            ("wFunc", wintypes.UINT),
            ("pFrom", wintypes.LPCWSTR),
            ("pTo", wintypes.LPCWSTR),
            ("fFlags", wintypes.WORD),
            ("fAnyOperationsAborted", wintypes.BOOL),
            ("hNameMappings", wintypes.LPVOID),
            ("lpszProgressTitle", wintypes.LPCWSTR),
        ]

    operation = SHFILEOPSTRUCTW()
    operation.wFunc = 3  # FO_DELETE
    operation.pFrom = f"{path}\0\0"
    operation.fFlags = 0x0040 | 0x0010 | 0x0400  # FOF_ALLOWUNDO | FOF_NOCONFIRMATION | FOF_NOERRORUI

    result = ctypes.windll.shell32.SHFileOperationW(ctypes.byref(operation))
    if result != 0 or operation.fAnyOperationsAborted:
        raise RuntimeError(f"Failed to move document to recycle bin: {result}")
